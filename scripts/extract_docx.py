import json
import re
from docx import Document

def extract_prompts(docx_path):
    doc = Document(docx_path)
    data = []
    
    current_category = "General"
    
    # regex for prompt starting with number: "1. ", "50. ", etc.
    prompt_regex = re.compile(r'^(\d+)\.\s*(.*)', re.DOTALL)
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        # Check if it's a prompt
        match = prompt_regex.match(text)
        if match:
            prompt_content = match.group(2).strip()
            data.append({
                "category": current_category,
                "content": prompt_content,
                "index": int(match.group(1))
            })
            continue
            
        # Check if it's an intro paragraph (suggests a new category)
        # "Aquí tienes X prompts para [CATEGORIA]..."
        if "tienes" in text.lower() and "prompt" in text.lower() and ("para" in text.lower() or "sobre" in text.lower()):
            # The category is usually in the PREVIOUS paragraph
            if i > 0:
                # Look back for the first non-empty paragraph
                for prev_idx in range(i-1, -1, -1):
                    prev_text = doc.paragraphs[prev_idx].text.strip()
                    if prev_text:
                        current_category = prev_text
                        break
        
        # If the paragraph is very short and no prompts follow immediately, it might be a section header
        # But our "intro paragraph" detection is more specific.
    
    return data

if __name__ == "__main__":
    docx_file = "1000 PROMPTS PRODUCTIVIDAD_HUBSPOT.docx"
    print(f"Extracting from {docx_file}...")
    results = extract_prompts(docx_file)
    
    print(f"Found {len(results)} prompts across multiple categories.")
    
    output_file = "scripts/raw_prompts.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    
    print(f"Saved to {output_file}")
